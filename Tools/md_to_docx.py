import re
from pathlib import Path

from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from langchain_core.tools import tool


# ---------- Low-level helpers ----------

def add_horizontal_rule(doc):
    """Insert a horizontal line using a paragraph bottom border (not a table)."""
    p = doc.add_paragraph()
    p_fmt = p.paragraph_format
    p_fmt.space_after = Pt(6)
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "999999")
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_page_break(doc):
    doc.add_page_break()


def add_inline_runs(paragraph, text):
    """
    Parse **bold** and *italic* (or _italic_) inside a line of text
    and add them as separate runs to the given paragraph.
    """
    # Split on **bold** and *italic*/_italic_ while keeping the delimiters
    tokens = re.split(r"(\*\*.+?\*\*|\*.+?\*|_.+?_)", text)

    for token in tokens:
        if not token:
            continue
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            run.bold = True
        elif (token.startswith("*") and token.endswith("*")) or \
             (token.startswith("_") and token.endswith("_")):
            run = paragraph.add_run(token[1:-1])
            run.italic = True
        else:
            paragraph.add_run(token)


def style_document(doc):
    """Set base fonts, margins, and default spacing for a professional look."""
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Give headings a consistent color/spacing
    for level, size in [(1, 20), (2, 16), (3, 13)]:
        try:
            h_style = doc.styles[f"Heading {level}"]
            h_style.font.size = Pt(size)
            h_style.font.bold = True
            h_style.font.name = "Calibri"
            h_style.paragraph_format.space_before = Pt(18 if level == 1 else 12)
            h_style.paragraph_format.space_after = Pt(6)
        except KeyError:
            pass


def add_title_page(doc, title):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(title)
    run.bold = True
    run.font.size = Pt(28)
    doc.add_paragraph()  # spacer
    add_page_break(doc)


# ---------- Table parsing ----------

def is_table_row(line):
    return line.strip().startswith("|") and line.strip().endswith("|")


def is_table_separator(line):
    # e.g. |---|---|---|  or  | :--- | :---: |
    stripped = line.strip().strip("|")
    cells = [c.strip() for c in stripped.split("|")]
    return all(re.fullmatch(r":?-+:?", c) for c in cells) and len(cells) > 0


def parse_table_row(line):
    cells = line.strip().strip("|").split("|")
    return [c.strip() for c in cells]


def add_table(doc, rows):
    """rows: list of lists of strings, first row = header."""
    if not rows:
        return
    n_cols = len(rows[0])
    table = doc.add_table(rows=0, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for r_idx, row_data in enumerate(rows):
        row_cells = table.add_row().cells
        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= n_cols:
                continue
            cell = row_cells[c_idx]
            cell.text = ""
            para = cell.paragraphs[0]
            add_inline_runs(para, cell_text)
            if r_idx == 0:
                for run in para.runs:
                    run.bold = True

    doc.add_paragraph()  # spacing after table


# ---------- Main Markdown parser ----------

def parse_markdown_to_docx(doc, markdown_text):
    lines = markdown_text.splitlines()
    i = 0
    n = len(lines)

    while i < n:
        line = lines[i]
        stripped = line.strip()

        # Blank line -> skip (spacing handled by styles)
        if stripped == "":
            i += 1
            continue

        # Horizontal rule: ---, ***, ___ (only if not a table separator context)
        if re.fullmatch(r"(-{3,}|\*{3,}|_{3,})", stripped):
            add_horizontal_rule(doc)
            i += 1
            continue

        # Page break marker: custom convention "<!--pagebreak-->"
        if stripped == "<!--pagebreak-->":
            add_page_break(doc)
            i += 1
            continue

        # Headings: #, ##, ###, ####
        heading_match = re.match(r"^(#{1,4})\s+(.*)", stripped)
        if heading_match:
            level = len(heading_match.group(1))
            text = heading_match.group(2).strip()
            heading = doc.add_heading(level=min(level, 4))
            add_inline_runs(heading, text)
            i += 1
            continue

        # Table: header row + separator row + data rows
        if is_table_row(line) and i + 1 < n and is_table_separator(lines[i + 1]):
            table_rows = [parse_table_row(line)]
            i += 2  # skip header + separator
            while i < n and is_table_row(lines[i]):
                table_rows.append(parse_table_row(lines[i]))
                i += 1
            add_table(doc, table_rows)
            continue

        # Numbered list: "1. item"
        numbered_match = re.match(r"^\d+\.\s+(.*)", stripped)
        if numbered_match:
            text = numbered_match.group(1)
            p = doc.add_paragraph(style="List Number")
            add_inline_runs(p, text)
            i += 1
            continue

        # Bullet list: "- item" or "* item"
        bullet_match = re.match(r"^[-*]\s+(.*)", stripped)
        if bullet_match:
            text = bullet_match.group(1)
            p = doc.add_paragraph(style="List Bullet")
            add_inline_runs(p, text)
            i += 1
            continue

        # Default: regular paragraph
        p = doc.add_paragraph()
        add_inline_runs(p, stripped)
        i += 1


# ---------- LangChain tool ----------

@tool
def md_to_docx(input_file: str, output_file: str, title: str = None) -> str:
    """
    Convert a Markdown (.md) file into a professionally formatted DOCX document.
    Supports headings (H1-H4), bold, italic, bullet lists, numbered lists,
    tables, and horizontal rules.

    Args:
        input_file: Path to the Markdown file to convert.
        output_file: Path where the DOCX should be saved.
        title: Optional title for a title page. If provided, a title page
               is added before the main content.
    """
    Path(output_file).parent.mkdir(parents=True, exist_ok=True)

    with open(input_file, "r", encoding="utf-8") as f:
        markdown_text = f.read()

    doc = Document()
    style_document(doc)

    if title:
        add_title_page(doc, title)

    parse_markdown_to_docx(doc, markdown_text)

    doc.save(output_file)

    return f"DOCX successfully saved to: {output_file}"